import os
import time
import base64
import yaml
import pandas as pd
import sys
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
# 核心修复：显式导入具体实现类，确保 PyInstaller 打包时不会漏掉此模块
try:
    import selenium.webdriver.chrome.webdriver
    import selenium.webdriver.chrome.service
except ImportError:
    pass

def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    curr_dir = os.path.dirname(os.path.abspath(__file__))
    if os.path.basename(curr_dir) == 'result':
        return os.path.dirname(curr_dir)
    return curr_dir

def load_config():
    """统一兼容打包后的路径定位逻辑"""
    base_dir = get_base_dir()
    config_path = os.path.join(base_dir, "config.yaml")
    if not os.path.exists(config_path):
        return {}
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except:
        return {}

def get_charts_via_js(html_path, output_dir, chart_mapping):
    """通过执行JS脚本直接从Canvas获取Base64图片"""
    print(f"-> 正在从看板提取图片: {html_path}")
    chrome_options = Options()
    chrome_options.add_argument('--headless')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--allow-file-access-from-files')
    chrome_options.add_argument('--disable-web-security')
    
    driver = webdriver.Chrome(options=chrome_options)
    driver.set_window_size(1200, 800)
    
    # 转换为 file:/// 协议以支持本地访问
    file_url = 'file:///' + os.path.abspath(html_path).replace('\\', '/')
    driver.get(file_url)
    
    # 等待渲染
    time.sleep(3) 
    
    images = {}
    for mark, div_id in chart_mapping.items():
        print(f"   正在提取图表 [{mark}]...")
        
        # 改进的 JS 脚本：增加存在性判定，防止 undefined 报错
        js_cmd = f"""
        var container = document.getElementById('{div_id}');
        if (!container) return "ERROR:ID_NOT_FOUND";
        var chart = echarts.getInstanceByDom(container);
        if (!chart) return "ERROR:CANVAS_NOT_READY";
        return chart.getDataURL({{
            type: 'png',
            pixelRatio: 3,
            backgroundColor: '#fff'
        }});
        """
        
        try:
            base64_str = None
            # 给予最大 5 秒的循环重试等待渲染完成
            for i in range(10): 
                driver.execute_script(f"var el = document.getElementById('{div_id}'); if(el) el.scrollIntoView();")
                result = driver.execute_script(js_cmd)
                
                if result and result.startswith("data:image/png;base64,"):
                    base64_str = result
                    break
                elif result == "ERROR:ID_NOT_FOUND":
                    print(f"      ⚠️ 警告: 容器 ID [{div_id}] 在当前 HTML 中不存在")
                    break
                
                time.sleep(0.5) # 每 0.5s 轮询一次
                
            if base64_str:
                img_data = base64.b64decode(base64_str.split(',')[1])
                temp_path = os.path.join(output_dir, f"temp_{div_id}.png")
                with open(temp_path, "wb") as f:
                    f.write(img_data)
                images[mark] = temp_path
                print("      ✅ 提取成功")
            else:
                print(f"      ❌ 提取失败: {mark} (超时或 Canvas 未加载)")
        except Exception as e:
            print(f"      ❌ 解析异常 {mark}: {e}")


    driver.quit()
    return images

def build_report():
    print("\n--- 开始图表植入流程 ---")
    config = load_config()
    if not config:
        print("💡 提示：未检测到有效配置，跳过图表集成。")
        return

    # 1. 动态常量初始化 (严格对齐 config)
    try:
        end_dt = pd.to_datetime(config['analysis_period']['end_date'])
        time_label = str(end_dt.month)
    except:
        time_label = "unknown"

    excel_path = config.get('file_config', {}).get('output_file', '')
    # 绝对尊重 config 定义的输出目录
    if excel_path:
        output_dir = os.path.dirname(excel_path)
    else:
        output_dir = os.path.join(get_base_dir(), 'result')

    # 路径自愈
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    HTML_FILE = os.path.join(output_dir, f"分析看板_{time_label}.html")
    TEMPLATE_FILE = os.path.join(output_dir, f"{time_label}月运营报告_gen.docx")
    OUTPUT_FILE = os.path.join(output_dir, f"{time_label}月运营报告_gen.docx")


    # 核心映射表：Word占位符(内部关键字) vs HTML Div ID
    # 恢复您最初的带星号样式，确保完全匹配
    CHART_MAPPING = {
        "**每月订单数量趋势**": "1b6e58a40b394e60857af14d5b7b8ec8",
        "**每月订单总金额**": "dc0de83a05b74081bb1f96cfa0bb85b3",
        "**订单总金额组成**": "eb93e670de6b4a5ba60bff478c1ec7bb",
        "**每月交易商品数量趋势**": "1e91645237624656a64ab2795d70e6b4",
        "**月订单总金额组成**": "f63ca06ccec641a5b862f5e3dd6f5de7"
    }

    if not os.path.exists(HTML_FILE):
        print(f"❌ 找不到 HTML 看板，无法植入图表: {HTML_FILE}")
        return

    # 2. 调用加固后的提取函数
    image_paths = get_charts_via_js(HTML_FILE, output_dir, CHART_MAPPING)

    if not image_paths:
        print("⚠️ 未提取到任何图表，保持原样。")
        return

    # 3. 操作 Word
    doc = Document(TEMPLATE_FILE)
    for para in doc.paragraphs:
        for mark, path in image_paths.items():
            # 兼容多种占位符形式
            m_variants = [f"{{{{{mark}}}}}", f"{{{mark}}}", mark]
            target = next((v for v in m_variants if v in para.text), None)
            
            if target:
                # 首先清空该段落已有文本以彻底抹除括号残留
                # 提示：如果段落有其他配文需保留，建议改为局部 replace
                para.text = para.text.replace(target, "")
                run = para.add_run()
                run.add_picture(path, width=Inches(6.0))
                # 注意：这里不再立即删除 path，防止多次引用导致后续段落找不到文件
                break

    # 4. 全部处理完后，一次性清理临时图元，并保存文档
    try:
        doc.save(OUTPUT_FILE)
        print(f"🚀 报告生成闭环：图表已成功集成至 {OUTPUT_FILE}")
        
        # 稳健地清理所有临时图片
        for path in set(image_paths.values()):
            try:
                if os.path.exists(path):
                    os.remove(path)
            except:
                pass
    except Exception as e:
        print(f"❌ 保存/清理阶段发生错误: {e}")


if __name__ == "__main__":
    build_report()