import os
import sys
import yaml
import pandas as pd
from docx import Document

def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def load_config():
    base_dir = get_base_dir()
    config_path = os.path.join(base_dir, "config.yaml")
    if not os.path.exists(config_path):
        config_path = "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

def get_paths():
    config = load_config()
    # 强制获取配置中的时间作为标签
    try:
        end_dt = pd.to_datetime(config['analysis_period']['end_date'])
        time_label = str(end_dt.month)
    except:
        time_label = "unknown"
    
    # 核心：绝对尊重 config 中的输出路径定义
    excel_path = config.get('file_config', {}).get('output_file', '')
    if not excel_path:
        # 兜底逻辑
        output_dir = os.path.join(get_base_dir(), 'result')
    else:
        output_dir = os.path.dirname(excel_path)
    
    # 如果配置的文件夹不存在，自动创建（确保“三月”这类文件夹能稳健存在）
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
        
    EXCEL_FILE = excel_path
    # 优先检测根目录下的原始模板
    TEMPLATE_FILE = os.path.join(get_base_dir(), "template", "运营报告模板.docx")
    if not os.path.exists(TEMPLATE_FILE):
        # 兜底到根目录直接找
        TEMPLATE_FILE = os.path.join(get_base_dir(), "运营报告模板.docx")
        
    OUTPUT_FILE = os.path.join(output_dir, f"{time_label}月运营报告_gen.docx")
    return EXCEL_FILE, TEMPLATE_FILE, OUTPUT_FILE


# --- 1. 基础工具函数 ---
def format_m(val):
    """金额格式化"""
    try:
        if pd.isna(val) or val == "" or val == "--": return "0.00"
        return f"{float(str(val).replace(',', '')):,.2f}"
    except:
        return "0.00"


def fill_table_data(table, df, mapping, max_rows=None, skip_col_idx=None):
    """通用表格填充"""
    if df is None or df.empty: return
    data = df.head(max_rows) if max_rows else df
    for i, (_, row) in enumerate(data.iterrows()):
        if i + 1 >= len(table.rows): break
        cells = table.rows[i + 1].cells
        for col_idx, excel_col in mapping.items():
            if col_idx < len(cells):
                if skip_col_idx is not None and col_idx == skip_col_idx:
                    cells[col_idx].text = ""
                else:
                    val = row.get(excel_col, "")
                    if any(x in str(excel_col) for x in ["金额", "总额", "单价", "计算"]):
                        cells[col_idx].text = format_m(val)
                    else:
                        cells[col_idx].text = str(val) if pd.notna(val) else ""


# --- 2. 主集成执行逻辑 ---
def main():
    print(f"🚀 开始全量集成化处理...")
    
    # 动态载入当前系统的时文路径和配置设定
    EXCEL_FILE, TEMPLATE_FILE, OUTPUT_FILE = get_paths()

    # A. 读取 Excel 数据
    sheets = pd.read_excel(EXCEL_FILE, sheet_name=None)
    df_core = sheets.get('核心提取数据')
    c_loaded = dict(zip(df_core['占位符'], df_core['填充值']))

    # B. 构建话术字典
    c = {
        "{{截止日期}}": "2026年03月31日",
        "{{产品名称}}": "阳光优采"
    }
    c.update(c_loaded)

    # C. 操作 Word
    doc = Document(TEMPLATE_FILE)

    # 1. 替换文字占位符 (高保真 Run 级版本，保留加粗/颜色/字体)
    # 组合所有段落和表格中的段落
    all_paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for cell in r.cells:
                all_paragraphs.extend(cell.paragraphs)

    for p in all_paragraphs:
        # 预检
        if not p.text or '{{' not in p.text:
            continue
            
        for k, v in c.items():
            if str(k) in p.text:
                # 核心：逐 Run 替换。docx 可能会把占位符切散，这里采用最稳健的“首个匹配 Run 写入+后续清空”逻辑
                # 先取出全文内容
                full_para_text = "".join(run.text for run in p.runs)
                if str(k) in full_para_text:
                    # 将替换后的内容注入
                    new_full_text = full_para_text.replace(str(k), str(v))
                    # 关键！我们保留第一个 run 的格式，并将新内容存入
                    if p.runs:
                        p.runs[0].text = new_full_text
                        # 清空该段落其余所有 runs，实现无缝替换且保留格式
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""

    # 2. 填充业务表格 (第3个表开始)
    ts = doc.tables
    fill_table_data(ts[2], sheets.get('采购企业汇总表'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '订单数量', 4: '订单总额（元）', 5: '首次订单日期',
                     6: '末次订单日期'})
    fill_table_data(ts[3], sheets.get('采购企业订单量TOP10'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '省市', 4: '订单数量', 5: '商品数量',
                     6: '订单金额（元）'}, 10)
    fill_table_data(ts[4], sheets.get('采购企业交易额TOP10'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '省市', 4: '订单数量', 5: '商品数量',
                     6: '订单金额（元）'}, 10)
    fill_table_data(ts[5], sheets.get('供应商汇总表'),
                    {0: '序号', 1: '供应商', 2: '供应商类型', 3: '订单数量', 4: '订单总额（元）', 5: '商品数量', 6: '专区名称'})
    fill_table_data(ts[7], sheets.get('供应商订单量TOP10'),
                    {0: '序号', 1: '单位名称', 2: '供应商类型', 3: '订单数量', 4: '商品数量', 5: '订单金额（元）'}, 10)
    fill_table_data(ts[8], sheets.get('供应商交易额TOP10'),
                    {0: '序号', 1: '单位名称', 2: '供应商类型', 3: '订单金额（元）', 4: '订单数量', 5: '商品数量'}, 10)
    fill_table_data(ts[9], sheets.get('商品销售数量TOP10'),
                    {0: '序号', 1: '商品名称', 2: '供应商', 3: '销售数量', 4: '销售总额_计算', 5: '平均单价',
                     6: '专区名称'}, 10)
    fill_table_data(ts[10], sheets.get('商品销售金额TOP10'),
                    {0: '序号', 1: '商品名称', 2: '供应商', 3: '销售数量', 4: '销售总额_计算', 5: '平均单价',
                     6: '专区名称'}, 10)

    # D. 保存生成的报告
    try:
        doc.save(OUTPUT_FILE)
        print(f"\n✨ 数据植入成功！报告：{OUTPUT_FILE}")
    except PermissionError:
        print(f"\n❌ 保存失败：文件 '{OUTPUT_FILE}' 可能正被其他程序打开（如 Word）。")
        print("💡 请关闭对应的 Word 文档后，重新运行脚本。")
    except Exception as e:
        print(f"\n❌ 保存最终报告时发生未知错误: {e}")


if __name__ == "__main__":
    main()