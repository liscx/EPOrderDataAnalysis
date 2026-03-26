import pandas as pd
from docx import Document
import os


# --- 1. 基础工具函数 ---
def format_m(val):
    """金额格式化：将数字转为 123,456.78 格式"""
    try:
        if pd.isna(val) or val == "" or val == "--": return "0.00"
        return f"{float(str(val).replace(',', '')):,.2f}"
    except:
        return "0.00"


def safe_split(val, expected_len):
    """严格按 | 切分 Excel 内容，长度不足则补 '0'"""
    if pd.isna(val) or val == "": return ["0"] * expected_len
    parts = str(val).split('|')
    if len(parts) < expected_len:
        parts.extend(["0"] * (expected_len - len(parts)))
    return parts


def fill_table_data(table, df, mapping, max_rows=None, skip_col_idx=None):
    """通用表格填充逻辑"""
    if df is None or df.empty: return
    data = df.head(max_rows) if max_rows else df

    # 从第2行开始填
    for i, (_, row) in enumerate(data.iterrows()):
        if i + 1 >= len(table.rows): break
        cells = table.rows[i + 1].cells
        for col_idx, excel_col in mapping.items():
            if col_idx < len(cells):
                # 如果是指定要空着的列（如供应商类型）
                if skip_col_idx is not None and col_idx == skip_col_idx:
                    cells[col_idx].text = ""
                else:
                    val = row.get(excel_col, "")
                    # 自动金额格式化
                    if any(x in str(excel_col) for x in ["金额", "总额", "单价", "计算"]):
                        cells[col_idx].text = format_m(val)
                    else:
                        cells[col_idx].text = str(val) if pd.notna(val) else ""


# --- 2. 主执行逻辑 ---
def generate_final_report(excel_path, template_path, output_path):
    print(f"🚀 开始生成最终运营报告...")
    sheets = pd.read_excel(excel_path, sheet_name=None)
    df_core = sheets.get('核心话术数据').set_index('维度')
    d = df_core.to_dict('index')

    # --- 映射字典构建 (已去掉本地供应商首次产生订单 & 新注册供应商话术) ---
    c = {
        "截止日期": "2026年03月31日",
        "产品名称": "阳光优采"
    }

    # 话术1: 概况
    v1 = safe_split(d['话术1']['数值'], 4)
    i1 = safe_split(d['话术1']['关联信息'], 4)
    c.update({
        "采购人数量": v1[0], "供应商数量": v1[1], "电商数量": v1[2], "本地供应商数量": v1[3],
        "订单数量": i1[0], "订单总额": format_m(i1[1]), "已完成订单数量": i1[2], "已完成订单总额": format_m(i1[3])
    })

    # 话术2: 累计专区
    i2 = safe_split(d['话术2']['关联信息'], 9)
    c.update({
        "产生订单专区数量": d['话术2']['数值'],
        "主要专区1": i2[1], "主要专区1订单": i2[2], "主要专区1占比": i2[3],
        "主要专区2": i2[4], "主要专区2订单": i2[5], "主要专区2占比": i2[6]
    })

    # 话术3: 本月表现
    v3 = safe_split(d['话术3']['数值'], 3)
    i3 = safe_split(d['话术3']['关联信息'], 10)
    cur_m, pre_m = float(i3[0]), float(i3[1])
    c.update({
        "月份": v3[0], "当月产生订单专区数量": v3[1], "当月订单数量": v3[2],
        "当月交易总额": format_m(cur_m), "上月交易总额": format_m(pre_m),
        "增长金额": format_m(cur_m - pre_m),
        "环比增长率": f"{((cur_m - pre_m) / pre_m * 100):.2f}%" if pre_m > 0 else "0.00%",
        "主要贡献专区": i3[2], "主要贡献专区订单": i3[3], "主要贡献专区总额": format_m(i3[4]),
        "次要贡献专区": i3[5], "次要贡献专区订单": i3[6], "次要贡献专区总额": format_m(i3[7])
    })

    # 话术4: 历史最值
    i4 = safe_split(d['话术4']['关联信息'], 8)
    c.update({
        "活跃采购人数量": d['话术4']['数值'],
        "最高订单数": i4[0], "最高订单采购人": i4[1], "最高订单总额": format_m(i4[2]), "最高订单专区": i4[3],
        "最高金额": format_m(i4[4]), "最高金额采购人": i4[5], "最高金额订单数": i4[6], "最高金额专区": i4[7]
    })

    # 话术5, 6, 9 (去掉了 8: 新注册供应商)
    c.update({
        "新注册采购人数量": d['话术5']['数值'], "新注册采购人专区": d['话术5']['关联信息'],
        "供应商总数": v1[1], "电商数量": v1[2], "本地供应商数量": v1[3],
        "商品总数": d['话术9']['数值'], "商品品类": d['话术9']['关联信息']
    })

    # --- 执行 Word 替换与填充 ---
    doc = Document(template_path)

    # 1. 文本占位符替换
    for target in list(doc.paragraphs) + [cell for t in doc.tables for r in t.rows for cell in r.cells]:
        for k, v in c.items():
            if f"{{{{{k}}}}}" in target.text:
                target.text = target.text.replace(f"{{{{{k}}}}}", str(v))

    # 2. 业务表格填充 (严格按要求从第3个表开始)
    ts = doc.tables

    # 表3: 采购汇总
    fill_table_data(ts[2], sheets.get('采购企业汇总表'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '订单数量', 4: '订单总额（元）', 5: '首次订单日期',
                     6: '末次订单日期'})
    # 表4: 采购订单TOP10
    fill_table_data(ts[3], sheets.get('采购企业订单量TOP10'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '省市', 4: '订单数量', 5: '商品数量',
                     6: '订单金额（元）'}, 10)
    # 表5: 采购交易额TOP10
    fill_table_data(ts[4], sheets.get('采购企业交易额TOP10'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '省市', 4: '订单数量', 5: '商品数量',
                     6: '订单金额（元）'}, 10)
    # 表6: 供应商汇总 (第3列索引2类型空着)
    fill_table_data(ts[5], sheets.get('供应商汇总表'),
                    {0: '序号', 1: '供应商', 3: '订单数量', 4: '订单总额（元）', 5: '商品数量', 6: '专区名称'},
                    skip_col_idx=2)
    # 表8: 供应商订单TOP10
    fill_table_data(ts[7], sheets.get('供应商订单量TOP10'),
                    {0: '序号', 1: '单位名称', 2: '供应商类型', 3: '订单数量', 4: '商品数量', 5: '订单金额（元）'}, 10)
    # 表9: 供应商交易额TOP10
    fill_table_data(ts[8], sheets.get('供应商交易额TOP10'),
                    {0: '序号', 1: '单位名称', 2: '供应商类型', 3: '订单金额（元）', 4: '订单数量', 5: '商品数量'}, 10)
    # 表10: 商品数量TOP10
    fill_table_data(ts[9], sheets.get('商品销售数量TOP10'),
                    {0: '序号', 1: '商品名称', 2: '供应商', 3: '销售数量', 4: '销售总额_计算', 5: '平均单价',
                     6: '专区名称'}, 10)
    # 表11: 商品金额TOP10
    fill_table_data(ts[10], sheets.get('商品销售金额TOP10'),
                    {0: '序号', 1: '商品名称', 2: '供应商', 3: '销售数量', 4: '销售总额_计算', 5: '平均单价',
                     6: '专区名称'}, 10)

    doc.save(output_path)
    print(f"✨ 任务完成！已剔除本地供应商及新注册供应商话术。文件：{output_path}")


if __name__ == "__main__":
    generate_final_report('数据清洗汇总结果.xlsx', '运营报告模板.docx', '3月运营报告_最终修正版.docx')