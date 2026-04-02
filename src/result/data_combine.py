import os
import sys
import yaml
import pandas as pd
from docx import Document

from ..core import get_base_dir, load_config

def get_paths():
    config = load_config()
    # 强制获取配置中的时间作为标签
    try:
        end_dt = pd.to_datetime(config['analysis_period']['end_date'])
        time_label = str(end_dt.month)
    except:
        time_label = "unknown"
    
    # 核心：绝对尊重 config 中的输出路径定义
    excel_path = config.get('file_config', {}).get('output_file', '')
    if not excel_path:
        # 兜底逻辑
        output_dir = os.path.join(get_base_dir(), 'result')
    else:
        output_dir = os.path.dirname(excel_path)
    
    # 如果配置的文件夹不存在，自动创建（确保“三月”这类文件夹能稳健存在）
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
        
    EXCEL_FILE = excel_path
    # 从核心常量库载入扫描路径规则
    from ..core.constants import TEMPLATE_SCAN_DIRS
    filename = "运营报告模板.docx"
    paths_to_try = [os.path.join(get_base_dir(), d, filename).replace("\\\\", "\\") for d in TEMPLATE_SCAN_DIRS]
    
    TEMPLATE_FILE = None
    for p in paths_to_try:
        if os.path.exists(p):
            TEMPLATE_FILE = p
            break
            
    if not TEMPLATE_FILE:
        # 如果还是没找到，默认回退到根目录路径（即使不存在也会在后续报错中抛出）
        TEMPLATE_FILE = os.path.join(get_base_dir(), "运营报告模板.docx")
        
    OUTPUT_FILE = os.path.join(output_dir, f"{time_label}月运营报告_gen.docx")
    return EXCEL_FILE, TEMPLATE_FILE, OUTPUT_FILE


# --- 1. 基础工具函数 ---
def format_m(val):
    """金额格式化"""
    try:
        if pd.isna(val) or val == "" or val == "--": return "0.00"
        return f"{float(str(val).replace(',', '')):,.2f}"
    except:
        return "0.00"


def fill_table_data(table, df, mapping, max_rows=None, skip_col_idx=None):
    """
    通用表格填充（自适应行数增强版）
    """
    if df is None or df.empty:
        # 如果没有数据，且表长超过1行（有表头），尝试清理冗余行
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)
        return

    data = df.head(max_rows) if max_rows else df
    target_row_count = len(data)
    
    # 1. 动态调整表格行数 (保持1个标题行 + target_row_count 条数据行)
    current_data_rows = len(table.rows) - 1
    
    # 如果模板行数多，删除多余行
    if current_data_rows > target_row_count:
        while len(table.rows) > target_row_count + 1:
            table._tbl.remove(table.rows[-1]._tr)
    
    # 如果模板行数少，动态新增并克隆样式 (从数据行第一行克隆样式)
    elif current_data_rows < target_row_count:
        for _ in range(target_row_count - current_data_rows):
            new_row = table.add_row()
            # 备注：docx.add_row 默认就会跟随上行样式，可满足基础需求

    # 2. 依次填充数据
    for i, (_, row) in enumerate(data.iterrows()):
        cells = table.rows[i + 1].cells
        for col_idx, excel_col in mapping.items():
            if col_idx < len(cells):
                if skip_col_idx is not None and col_idx == skip_col_idx:
                    cells[col_idx].text = ""
                else:
                    val = row.get(excel_col, "")
                    if any(x in str(excel_col) for x in ["金额", "总额", "单价", "计算"]):
                        cells[col_idx].text = format_m(val)
                    else:
                        cells[col_idx].text = str(val) if pd.notna(val) else ""



# --- 2. 主集成执行逻辑 ---
def main():
    print(f"🚀 开始全量集成化处理...")
    
    # 动态载入当前系统的时文路径和配置设定
    EXCEL_FILE, TEMPLATE_FILE, OUTPUT_FILE = get_paths()

    # A. 读取 Excel 数据
    sheets = pd.read_excel(EXCEL_FILE, sheet_name=None)
    df_core = sheets.get('核心提取数据')
    c_loaded = dict(zip(df_core['占位符'], df_core['填充值']))

    # B. 构建话术字典
    c = {
        "{{截止日期}}": "2026年03月31日",
        "{{产品名称}}": "阳光优采"
    }
    c.update(c_loaded)

    # C. 操作 Word
    doc = Document(TEMPLATE_FILE)

    # 1. 替换文字占位符 (高保真 Run 级版本，保留加粗/颜色/字体)
    # 组合所有段落和表格中的段落
    all_paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for cell in r.cells:
                all_paragraphs.extend(cell.paragraphs)

    for p in all_paragraphs:
        # 预检
        if not p.text or '{{' not in p.text:
            continue
            
        for k, v in c.items():
            if str(k) in p.text:
                # 核心：逐 Run 替换。docx 可能会把占位符切散，这里采用最稳健的“首个匹配 Run 写入+后续清空”逻辑
                # 先取出全文内容
                full_para_text = "".join(run.text for run in p.runs)
                if str(k) in full_para_text:
                    # 将替换后的内容注入
                    new_full_text = full_para_text.replace(str(k), str(v))
                    # 关键！我们保留第一个 run 的格式，并将新内容存入
                    if p.runs:
                        p.runs[0].text = new_full_text
                        # 清空该段落其余所有 runs，实现无缝替换且保留格式
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""

    # 2. 填充业务表格 (第3个表开始)
    ts = doc.tables
    fill_table_data(ts[2], sheets.get('采购企业汇总表'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '订单数量', 4: '订单总额（元）', 5: '首次订单日期',
                     6: '末次订单日期'})
    fill_table_data(ts[3], sheets.get('采购企业订单量TOP10'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '省市', 4: '订单数量', 5: '商品数量',
                     6: '订单金额（元）'}, 10)
    # 【已修正】Table 4 (索引ts[4]): 序号 | 采购企业 | 专区名称 | 省市 | 订单金额（元） | 订单数量 | 商品数量
    fill_table_data(ts[4], sheets.get('采购企业交易额TOP10'),
                    {0: '序号', 1: '采购企业', 2: '专区名称', 3: '省市', 4: '订单金额（元）', 5: '订单数量',
                     6: '商品数量'}, 10)
    fill_table_data(ts[5], sheets.get('供应商汇总表'),
                    {0: '序号', 1: '供应商', 2: '供应商类型', 3: '订单数量', 4: '订单总额（元）', 5: '商品数量', 6: '专区名称'})
    fill_table_data(ts[7], sheets.get('供应商订单量TOP10'),
                    {0: '序号', 1: '单位名称', 2: '供应商类型', 3: '订单数量', 4: '商品数量', 5: '订单金额（元）'}, 10)
    fill_table_data(ts[8], sheets.get('供应商交易额TOP10'),
                    {0: '序号', 1: '单位名称', 2: '供应商类型', 3: '订单金额（元）', 4: '订单数量', 5: '商品数量'}, 10)
    # 【已修正】Table 9 (索引ts[9]): 序号 | 商品名称 | 供应商 | 销售数量 | 单价（元） | 销售总额（元） | 专区名称
    fill_table_data(ts[9], sheets.get('商品销售数量TOP10'),
                    {0: '序号', 1: '商品名称', 2: '供应商', 3: '销售数量', 4: '平均单价', 5: '销售总额_计算',
                     6: '专区名称'}, 10)
    # 【已修正】Table 10 (索引ts[10]): 序号 | 商品名称 | 供应商 | 销售总额（元） | 销售数量 | 单价（元） | 专区
    fill_table_data(ts[10], sheets.get('商品销售金额TOP10'),
                    {0: '序号', 1: '商品名称', 2: '供应商', 3: '销售总额_计算', 4: '销售数量', 5: '平均单价',
                     6: '专区名称'}, 10)


    # 【升级：绝对全等锁定表格】严格检查表头序列：['序号', '单位名称', '专区名称', '入库日期']
    target_sheet = sheets.get('供应商信息提取')
    if target_sheet is not None and not target_sheet.empty:
        found_table = False
        target_headers = ['序号', '单位名称', '专区名称', '入库日期']
        for table in ts:
            try:
                # 获取第一行所有单元格内容
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                # 只有当且仅当每一个表头完全一致时才填入
                if len(header_cells) >= 4 and header_cells[:4] == target_headers:
                    fill_table_data(table, target_sheet,
                                    {0: '序号', 1: '单位名称', 2: '专区名称', 3: '入库日期'})
                    found_table = True
                    print(">>> 已成功在 Word 中精准锁定并填充 [新注册供应商清单] 表格。")
                    break
            except:
                continue
        if not found_table:
            print(f"警告：未在 Word 模板中检测到全等表头 {target_headers} 的表格，跳过明细填充。")

    # D. 保存生成的报告
    try:
        doc.save(OUTPUT_FILE)


        print(f"\n✨ 数据植入成功！报告：{OUTPUT_FILE}")
    except PermissionError:
        print(f"\n❌ 保存失败：文件 '{OUTPUT_FILE}' 可能正被其他程序打开（如 Word）。")
        print("💡 请关闭对应的 Word 文档后，重新运行脚本。")
    except Exception as e:
        print(f"\n❌ 保存最终报告时发生未知错误: {e}")


if __name__ == "__main__":
    main()